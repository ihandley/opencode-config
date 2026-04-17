from pathlib import Path
from typing import Any, Dict
import re

from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt


_PROJECT_ROOT = Path.home() / "code" / "github" / "opencode"
BASE_DIR = _PROJECT_ROOT / "data" / "job-coach"


# ---------------------------------------------------------------------------
# Inline markdown → rich-text runs
# ---------------------------------------------------------------------------

# Order matters: bold must match before italic so ** is consumed first.
# Each group set:  (full_match, captured_text)
_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*)"           # group 1,2  — bold+italic
    r"|(\*\*(.+?)\*\*)"              # group 3,4  — bold
    r"|(__(.+?)__)"                   # group 5,6  — bold (underscores)
    r"|(\*(.+?)\*)"                   # group 7,8  — italic
    r"|(_(.+?)_)"                     # group 9,10 — italic (underscores)
    r"|(`(.+?)`)"                     # group 11,12 — inline code
    r"|(\[([^\]]+)\]\(([^)]+)\))"     # group 13,14,15 — link [text](url)
)


def _add_rich_text(paragraph: Paragraph, text: str) -> None:
    """Parse inline markdown and add formatted runs to *paragraph*."""
    last = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match
        if m.start() > last:
            paragraph.add_run(text[last : m.start()])

        if m.group(2):       # ***bold italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(4):     # **bold**
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif m.group(6):     # __bold__
            run = paragraph.add_run(m.group(6))
            run.bold = True
        elif m.group(8):     # *italic*
            run = paragraph.add_run(m.group(8))
            run.italic = True
        elif m.group(10):    # _italic_
            run = paragraph.add_run(m.group(10))
            run.italic = True
        elif m.group(12):    # `code`
            run = paragraph.add_run(m.group(12))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif m.group(14):    # [text](url) — render text only (no hyperlink API)
            paragraph.add_run(m.group(14))

        last = m.end()

    # Remaining plain text after the last match
    if last < len(text):
        paragraph.add_run(text[last:])


# ---------------------------------------------------------------------------
# Line-type classification helpers
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")
_BULLET_RE = re.compile(r"^(\s*)([-*])\s+(.*)")
_NUMBERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)")
_HR_RE = re.compile(r"^[\s]*([-*_])\s*\1\s*\1[\s\-*_]*$")


def _is_metadata_line(line: str) -> bool:
    stripped = line.strip()
    if _HR_RE.match(stripped):
        return True
    return stripped.lower().startswith("resume tailored for")


# ---------------------------------------------------------------------------
# Public export function
# ---------------------------------------------------------------------------

def export_resume_version_docx(filename: str, company: str = None) -> Dict[str, Any]:
    if not filename:
        return {
            "ok": False,
            "input_path": "",
            "output_path": "",
            "error": "No filename provided",
        }

    # Determine input and output paths
    if company:
        company_slug = company.lower().strip()
        company_slug = re.sub(r"[^a-z0-9]+", "-", company_slug).strip("-")
        input_path = BASE_DIR / company_slug / filename
        output_dir = BASE_DIR / company_slug
    else:
        # Try to find in any company directory
        input_path = None
        output_dir = None
        for company_dir in BASE_DIR.iterdir():
            if company_dir.is_dir():
                candidate = company_dir / filename
                if candidate.is_file():
                    input_path = candidate
                    output_dir = company_dir
                    break
        if input_path is None:
            return {
                "ok": False,
                "input_path": "",
                "output_path": "",
                "error": f"File not found: {filename}",
            }

    if not input_path.is_file():
        return {
            "ok": False,
            "input_path": str(input_path),
            "output_path": "",
            "error": f"File not found: {input_path}",
        }

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{input_path.stem}.docx"

    try:
        text = input_path.read_text(encoding="utf-8")
        lines = text.splitlines()

        doc = Document()

        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()

            # Skip metadata / horizontal rules
            if _is_metadata_line(stripped):
                continue

            # Skip blank lines
            if not stripped:
                continue

            # --- Headings (#{1,6}) ---
            hm = _HEADING_RE.match(stripped)
            if hm:
                level = min(len(hm.group(1)), 4)  # DOCX supports levels 1-4
                heading = doc.add_heading(level=level)
                _add_rich_text(heading, hm.group(2).strip())
                continue

            # --- Bullet lists (- or *) with indentation ---
            bm = _BULLET_RE.match(line)
            if bm:
                indent_depth = len(bm.group(1))
                # Use "List Bullet" for top-level, "List Bullet 2" for nested
                style = "List Bullet 2" if indent_depth >= 2 else "List Bullet"
                para = doc.add_paragraph(style=style)
                _add_rich_text(para, bm.group(3).strip())
                continue

            # --- Numbered lists (1. text) ---
            nm = _NUMBERED_RE.match(line)
            if nm:
                indent_depth = len(nm.group(1))
                style = "List Number 2" if indent_depth >= 2 else "List Number"
                para = doc.add_paragraph(style=style)
                _add_rich_text(para, nm.group(3).strip())
                continue

            # --- Normal paragraph with rich-text ---
            para = doc.add_paragraph()
            _add_rich_text(para, stripped)

        doc.save(str(output_path))
        return {
            "ok": True,
            "input_path": str(input_path),
            "output_path": str(output_path),
            "error": "",
        }
    except Exception as e:
        return {
            "ok": False,
            "input_path": str(input_path),
            "output_path": str(output_path) if output_path else "",
            "error": str(e),
        }
